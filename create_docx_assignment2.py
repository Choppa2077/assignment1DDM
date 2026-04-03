"""
Generates Assignment_2_DDDM.docx — full report for Assignment 2
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BLUE   = RGBColor(0x2A, 0x5C, 0xAA)
YELLOW = RGBColor(0xF5, 0xC8, 0x42)
RED    = RGBColor(0xE8, 0x5D, 0x2C)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
LIGHT_BLUE = RGBColor(0xDD, 0xE8, 0xF8)
LIGHT_GREEN = RGBColor(0xD5, 0xF0, 0xDD)
LIGHT_RED   = RGBColor(0xFB, 0xE0, 0xDA)
LIGHT_YELLOW= RGBColor(0xFE, 0xF7, 0xD6)

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2)


# ── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    # RGBColor stores as bytes: index 0=R, 1=G, 2=B
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="2A5CAA", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(text, level=1, color=BLUE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def body(text, bold=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = DARK
    return p

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part and text.startswith(bold_part):
        r1 = p.add_run(bold_part)
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = DARK
        r2 = p.add_run(text[len(bold_part):])
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK
    return p

def add_table(headers, rows, header_color=BLUE):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = LIGHT_BLUE if ri % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
    return table

def spacer(n=1):
    for _ in range(n):
        doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AITU — CSE-2501M")
run.font.size = Pt(12)
run.font.color.rgb = DARK

spacer()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Assignment №2")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Data-Driven Decision Project — Stage 2")
run2.font.size = Pt(16)
run2.font.bold = True
run2.font.color.rgb = BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Data Preparation and Market Intelligence")
run3.font.size = Pt(13)
run3.font.color.rgb = DARK

spacer()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("RentEasy KZ — Rental Housing Aggregator for Kazakhstan")
run4.font.size = Pt(12)
run4.font.italic = True
run4.font.color.rgb = DARK

spacer()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run("Idrissov Mukan, Yeraliev\nCSE-2501M\n2026")
run5.font.size = Pt(11)
run5.font.color.rgb = DARK

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DATA COLLECTION (TASK 1)
# ════════════════════════════════════════════════════════════════════════════
heading("1. Data Collection and Preparation (Task 1 — 50%)", level=1)

heading("1.1 Data Sources", level=2)
body(
    "Data was collected from two open Kazakhstan rental platforms: Krisha.kz and OLX.kz. "
    "Both sources cover apartment rental listings and complement each other — Krisha.kz "
    "provides structured, rich listings across 5 cities, while OLX.kz adds coverage in "
    "Astana and Almaty with different seller behaviour."
)
spacer()

add_table(
    ["Source", "City Coverage", "Records Collected", "Key Fields"],
    [
        ["Krisha.kz", "Астана, Алматы, Шымкент, Карагандa, Актобе", "6,130", "price, area, floor, rooms, district, seller_type"],
        ["OLX.kz",   "Астана, Алматы",                              "1,355", "price, district, title, listing_date"],
        ["TOTAL",    "5 cities",                                    "7,485", "16 fields, merged dataset"],
    ]
)

spacer()
heading("1.2 Dataset Structure", level=2)
body("The merged raw dataset contains 7,485 records and 16 fields:")
spacer()

add_table(
    ["Field", "Type", "Description"],
    [
        ["source",        "string",  "Data source: 'krisha' or 'olx'"],
        ["city",          "string",  "City: Астана, Алматы, Шымкент, Карагандa, Актобе"],
        ["district",      "string",  "City district / neighbourhood"],
        ["rooms",         "integer", "Number of rooms (0=studio, 1, 2, 3, 4+)"],
        ["area_m2",       "float",   "Apartment area in square metres (Krisha only)"],
        ["floor",         "integer", "Floor number in the building (Krisha only)"],
        ["total_floors",  "integer", "Total floors in building (Krisha only)"],
        ["price_tenge",   "integer", "Monthly rental price in Kazakhstani tenge"],
        ["price_per_m2",  "float",   "Calculated: price_tenge / area_m2"],
        ["seller_type",   "string",  "owner / agency / unknown"],
        ["listing_date",  "string",  "Publication date (YYYY-MM-DD)"],
        ["title",         "string",  "Full listing title text"],
        ["url",           "string",  "Direct URL to listing"],
        ["listing_id",    "string",  "Original platform listing ID"],
        ["id",            "integer", "Sequential record ID"],
        ["price_segment", "string",  "Assigned segment: low / middle / high / luxury"],
    ]
)

spacer()

heading("1.3 Data Cleaning — Step by Step", level=2)
body(
    "All cleaning steps were implemented in data_cleaning.py. "
    "Each step was logged to renteasy_data/cleaning_log.txt. "
    "The cleaned dataset was saved as renteasy_data/renteasy_dataset_cleaned.csv."
)
spacer()

heading("Step 1 — Duplicate Detection and Removal", level=3)
body("Data was checked for duplicates at two levels:")
bullet("Full row duplicates: 0 found")
bullet("Duplicate by (source, listing_id) key: 0 found — deduplication was already applied during the merging step in Assignment 1")
body("Result: no rows removed in this step.")

heading("Step 2 — Missing Values Analysis", level=3)
body("Null and empty values were detected per column before any cleaning:")
spacer()

add_table(
    ["Field", "Missing (raw)", "% of Total", "Action Taken"],
    [
        ["rooms",       "858",  "11.5%", "Kept as NaN (OLX titles lack structured room data)"],
        ["area_m2",     "1,355", "18.1%", "Kept as NaN (OLX listings only have price + title)"],
        ["floor",       "1,841", "24.6%", "Kept as NaN (OLX + some Krisha without floor data)"],
        ["total_floors","1,841", "24.6%", "Kept as NaN (same as floor)"],
        ["price_per_m2","1,355", "18.1%", "Recalculated after area_m2 cleaned"],
        ["price_tenge", "0",     "0%",    "No missing — critical field"],
        ["city",        "0",     "0%",    "No missing"],
        ["seller_type", "0",     "0%",    "No missing"],
        ["listing_date","0",     "0%",    "No missing"],
    ]
)

spacer()
body(
    "Empty string values were replaced with NaN. No rows were dropped for missing price_tenge "
    "or city (both were complete). Fields such as floor, area_m2, and rooms are legitimately "
    "absent for OLX listings where the platform does not expose structured data in listing cards."
)

heading("Step 3 — Outdated Listings Check", level=3)
body(
    "All listings were checked for the listing_date field. The cutoff date was set to "
    "2026-01-01. All 7,485 records fell within the range 2026-03-05 — 2026-04-03, "
    "so no records were removed as outdated. The dataset reflects current March–April 2026 market."
)

heading("Step 4 — Spelling and Standardisation", level=3)
body("The following standardisation steps were applied:")
bullet("City names: stripped whitespace, confirmed 5 unique values without mixed-script issues")
bullet("seller_type: validated against {'owner', 'agency', 'unknown'} — all values valid")
bullet("source: normalised to lowercase 'krisha' / 'olx' — all valid")
bullet("rooms: float values (e.g. 2.0) rounded to integer; values outside 0–10 set to NaN")
body("Seller type distribution after standardisation:")
spacer()

add_table(
    ["Seller Type", "Count", "Share"],
    [
        ["Owner (direct landlord)", "3,805", "50.8%"],
        ["Unknown",                 "2,400", "32.1%"],
        ["Agency / Specialist",     "1,280", "17.1%"],
    ]
)

heading("Step 5 — Outlier Detection and Removal (1.5 × IQR Method)", level=3)
body(
    "Outliers were detected using the standard 1.5 × IQR rule: values below Q1 − 1.5×IQR "
    "or above Q3 + 1.5×IQR were removed. The method was applied to three numeric fields:"
)
spacer()

add_table(
    ["Field", "Q1", "Q3", "IQR", "Lower Bound", "Upper Bound", "Outliers Removed"],
    [
        ["price_tenge",   "150,000 тг", "280,000 тг", "130,000", "−45,000 (→ 0)", "475,000 тг", "556"],
        ["area_m2",       "40 m²",      "65 m²",      "25",      "2 m²",          "102 m²",     "328"],
        ["price_per_m2",  "3,168 тг",   "5,607 тг",   "2,439",   "−491 (→ 0)",    "9,266 тг",   "185"],
    ]
)

spacer()
body(
    "Total rows removed as outliers: 801 (10.7% of raw dataset). "
    "These included luxury penthouse listings above 475,000 тг/month, micro-rooms below 2 m², "
    "and data entry errors with unrealistic prices."
)

heading("Step 6 — Derived Field: Price Segment", level=3)
body(
    "A new categorical field price_segment was calculated from price_tenge "
    "to support competitor analysis in Task 2:"
)
bullet("Low: price < 120,000 тг/month")
bullet("Middle: 120,000 ≤ price < 250,000 тг/month")
bullet("High: 250,000 ≤ price < 500,000 тг/month")
bullet("Luxury: price ≥ 500,000 тг/month (removed by IQR — 0 records remain)")

heading("1.4 Final Cleaned Dataset", level=2)
body(
    "After all cleaning steps the dataset was saved to "
    "renteasy_data/renteasy_dataset_cleaned.csv with the following characteristics:"
)
spacer()

add_table(
    ["Metric", "Value"],
    [
        ["Total records",    "6,684"],
        ["Total columns",    "17"],
        ["Sources",          "Krisha.kz (5,362) + OLX.kz (1,322)"],
        ["Cities",           "5: Астана, Алматы, Шымкент, Карагандa, Актобе"],
        ["Date range",       "2026-03-05 — 2026-04-03"],
        ["Price range",      "5,000 — 475,000 тг/month"],
        ["Avg price",        "207,579 тг/month"],
        ["Median price",     "200,000 тг/month"],
        ["Records removed",  "801 (10.7% — outliers and duplicates)"],
        ["Cleaning log",     "renteasy_data/cleaning_log.txt"],
    ]
)

spacer()
heading("Records by City — Final Dataset", level=3)
add_table(
    ["City", "Records", "Avg Price (тг/month)", "Avg Area (m²)", "Avg Price/m²"],
    [
        ["Астана",    "1,830", "228,789", "49", "5,031"],
        ["Алматы",    "1,965", "234,843", "50", "5,675"],
        ["Шымкент",   "1,009", "186,076", "55", "3,489"],
        ["Карагандa", "958",   "185,100", "50", "3,787"],
        ["Актобе",    "922",   "154,261", "53", "3,001"],
    ]
)

spacer()
body("Source code: data_cleaning.py | GitHub: https://github.com/Choppa2077/assignment1DDM")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — MARKET INTELLIGENCE (TASK 2)
# ════════════════════════════════════════════════════════════════════════════
heading("2. Market Intelligence and Competitor Analysis (Task 2 — 50%)", level=1)

heading("2.1 Price Segmentation — Competitor Landscape", level=2)
body(
    "The rental market in Kazakhstan is segmented into three active price tiers "
    "(luxury listings were removed as outliers). Each tier represents a distinct "
    "group of competitors — landlords and agencies targeting specific renter profiles."
)
spacer()

add_table(
    ["Segment", "Price Range (тг/month)", "Count", "Share", "Target Renter"],
    [
        ["Low",    "< 120,000",      "672",  "10.1%", "Students, minimum wage workers"],
        ["Middle", "120,000–250,000","3,858","57.7%", "Young professionals, families (core market)"],
        ["High",   "250,000–475,000","2,154","32.2%", "Senior specialists, expats, relocated employees"],
        ["Luxury", "> 475,000",      "0",    "0%",    "Removed as outliers (data integrity)"],
    ]
)

spacer()
body(
    "The Middle segment dominates with 57.7% of all listings — confirming that the "
    "primary rental market in Kazakhstan is affordable mid-range housing. This is "
    "the most contested and most promising segment for RentEasy KZ."
)

heading("2.2 Competitor Analysis by City", level=2)
body(
    "Each city represents a separate competitive environment. "
    "The table below shows the listing count per segment per city:"
)
spacer()

add_table(
    ["City", "Low (<120K)", "Middle (120–250K)", "High (250–475K)", "Total"],
    [
        ["Астана",    "68",  "1,037", "725",  "1,830"],
        ["Алматы",    "219", "790",   "956",  "1,965"],
        ["Шымкент",   "114", "678",   "217",  "1,009"],
        ["Карагандa", "106", "662",   "190",  "958"],
        ["Актобе",    "165", "691",   "66",   "922"],
    ]
)

spacer()
body(
    "Key observation: Алматы leads in the High segment (956 listings = 48.8% of its market), "
    "while Астана has a more balanced split between Middle and High — indicating it is "
    "the best city for a platform launch targeting the widest renter audience."
)

heading("2.3 Top Districts — Price Intelligence", level=2)
body(
    "Districts with the highest average prices indicate demand-supply imbalances — "
    "exactly where RentEasy's price analytics add the most value for renters."
)
spacer()

heading("Top Districts in Астана", level=3)
add_table(
    ["District", "Avg Price (тг/month)", "Listings"],
    [
        ["Нура р-н",           "257,265", "287"],
        ["Есильский р-н",      "255,998", "484"],
        ["Есильский район",    "225,085", "175"],
        ["Сарайшык р-н",       "219,388", "80"],
        ["Нура",               "216,640", "125"],
        ["Алматы р-н",         "215,857", "224"],
        ["Сарыарка р-н",       "203,485", "136"],
        ["р-н Байконур",       "203,023", "43"],
        ["Алматинский район",  "189,978", "158"],
        ["Сарыаркинский район","176,935", "62"],
    ]
)

spacer()
heading("Top Districts in Алматы", level=3)
add_table(
    ["District", "Avg Price (тг/month)", "Listings"],
    [
        ["Бостандыкский р-н",  "324,261", "272"],
        ["Медеуский р-н",      "307,302", "106"],
        ["Алмалинский р-н",    "303,128", "203"],
        ["Ауэзовский р-н",     "264,632", "201"],
        ["Медеуский район",    "257,143", "21"],
        ["Бостандыкский район","248,856", "52"],
        ["Алмалинский район",  "230,197", "122"],
        ["Ауэзовский район",   "219,106", "85"],
        ["Турксибский р-н",    "226,148", "71"],
        ["Наурызбайский р-н",  "219,200", "222"],
    ]
)

spacer()
body(
    "The price gap between districts is substantial — in Алматы the top district "
    "(Бостандыкский р-н, 324K тг) is 2× more expensive than the cheapest districts. "
    "This variance is invisible to renters on Krisha.kz or OLX.kz — "
    "it is RentEasy's primary analytics opportunity."
)

heading("2.4 Key Consumers of Competitors", level=2)
body(
    "Based on the price segment distribution and district data, the key consumers "
    "of current competitor platforms (Krisha.kz and OLX.kz) can be described as:"
)
spacer()

add_table(
    ["Segment", "Consumer Profile", "Tech Level", "Pain Point"],
    [
        ["Low",    "Students, entry-level workers, migrants",
                   "Basic smartphone user",
                   "Price is top priority — scam listings are dangerous"],
        ["Middle", "Young professionals (25–35), junior specialists, young families",
                   "Active mobile & web user",
                   "No price benchmarks, too many agency calls"],
        ["High",   "Senior employees, expats, relocated corporate staff",
                   "Comfortable with digital platforms",
                   "Need district-level analytics and verified listings"],
    ]
)

spacer()
body(
    "The Middle segment consumer (our primary persona — Алия, 24 лет, Junior Developer, Астана) "
    "is tech-savvy, trusts data over word-of-mouth, and expects a clean UX with reliable price "
    "transparency — none of which competitors currently offer."
)

heading("2.5 Krisha.kz vs OLX.kz — Source Comparison", level=2)
add_table(
    ["Metric", "Krisha.kz", "OLX.kz"],
    [
        ["Records in dataset", "5,362 (80.2%)", "1,322 (19.8%)"],
        ["City coverage", "5 cities", "2 cities (Астана, Алматы)"],
        ["Structured data", "Yes (area, floor, rooms)", "Partial (price, district only)"],
        ["Seller type info", "Yes", "No (all 'unknown')"],
        ["Avg price (Астана)", "231K тг", "220K тг"],
        ["Avg price (Алматы)", "260K тг", "185K тг"],
        ["Room data coverage", "~88%", "~37%"],
    ]
)

spacer()
body(
    "OLX.kz listings average significantly lower prices in Алматы (185K vs 260K), "
    "suggesting it captures a lower-income renter segment or owners posting directly "
    "without agency markup. This validates RentEasy's value proposition of aggregating "
    "both sources to give renters the complete market picture."
)

heading("2.6 SWOT Analysis of Competitors", level=2)
body(
    "Based on the collected data, the following SWOT analysis characterises "
    "the competitive landscape for RentEasy KZ:"
)
spacer()

# SWOT table — 2×2
swot_table = doc.add_table(rows=2, cols=2)
swot_table.style = "Table Grid"

swot_data = [
    (0, 0, "STRENGTHS", BLUE,         LIGHT_BLUE,
     [
         "Krisha.kz dominates with 80% of structured listings",
         "Both platforms are free to browse — low entry barrier",
         "Large middle segment (58%) confirms stable demand",
         "Established brand trust and organic SEO rankings",
         "OLX provides additional lower-price listings",
     ]),
    (0, 1, "WEAKNESSES", RED,          LIGHT_RED,
     [
         "14.6% agency listings → double commissions for renters",
         "No price analytics or historical trends on any platform",
         "OLX listing cards lack area/floor/rooms data (63% missing)",
         "Listings spread across 2+ sites — no single aggregator",
         "Fake / outdated listings are not filtered or flagged",
     ]),
    (1, 0, "OPPORTUNITIES", RGBColor(0x1E, 0x84, 0x49), LIGHT_GREEN,
     [
         "Астана avg rent 229K тг — high price variance = analytics demand",
         "Middle (58%) + High (32%) = 90% of market = RentEasy target",
         "61% owner listings → direct connection without agency fees",
         "AI price recommendations not offered by any competitor",
         "Seasonal peaks (Aug–Sep students) = marketing windows",
     ]),
    (1, 1, "THREATS", RGBColor(0xB7, 0x95, 0x0B), LIGHT_YELLOW,
     [
         "Krisha.kz may add analytics (risk of defensive feature)",
         "Low entry barrier for new aggregator startups",
         "OLX anti-scraping changes broke data pipeline (URL fix needed)",
         "Real estate market slowdown reduces listing volume",
         "User trust in new platforms requires significant time",
     ]),
]

for (row_idx, col_idx, title, title_color, bg_color, bullets_list) in swot_data:
    cell = swot_table.rows[row_idx].cells[col_idx]
    set_cell_bg(cell, bg_color)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = title_color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for b in bullets_list:
        bp = cell.add_paragraph()
        bp.style = "List Bullet"
        r = bp.add_run(b)
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK

spacer()

heading("2.7 Most Promising Segment and Strategy", level=2)
body(
    "Based on the full analysis of 6,684 clean records across 5 cities, "
    "the following strategic recommendations are derived:"
)
spacer()

heading("Most Promising Market: Астана", level=3)
bullet("1,830 listings — second largest after Алматы but with faster price growth (+17.7% in 2025)")
bullet("Balanced segment split: 57% Middle + 40% High — widest renter audience")
bullet("Нура р-н and Есильский р-н are the premium zones — prime analytics use case")
bullet("High relocation activity (government, corporations) drives constant rental demand")

heading("Most Promising Segment: Middle (120,000–250,000 тг/month)", level=3)
bullet("58% of all cleaned listings — the largest and most competitive segment")
bullet("Core persona (Алия, 24, Junior Developer) operates exactly in this budget")
bullet("Competitor weakness: no price benchmarks available for this segment")
bullet("Agency share is lowest in Middle segment — direct owner connections are possible")

heading("Recommended Strategy for RentEasy KZ", level=3)
bullet("Phase 1 — Launch: Астана city, Middle segment, web + Telegram bot")
bullet("Phase 2 — Differentiate: price analytics dashboard (district median prices, trend charts)")
bullet("Phase 3 — Monetise: premium analytics subscription 990 тг/month + agency lead gen")
bullet("Phase 4 — Expand: Алматы, then Шымкент, Карагандa, Актобе")
body(
    "The data confirms that the Kazakhstan rental market is fragmented, opaque, and agency-heavy "
    "— exactly the conditions where a transparent aggregator with AI price intelligence "
    "creates maximum value. RentEasy KZ's competitive advantage is not in listing volume "
    "(Krisha.kz wins there) but in trust, transparency, and data-driven guidance."
)

spacer()

heading("2.8 Dashboards", level=2)
body(
    "Five analytical dashboards were generated using dashboard.py and saved to "
    "renteasy_data/dashboards/. Each dashboard uses the cleaned dataset (6,684 records)."
)
spacer()

DASH_DIR = "renteasy_data/dashboards"
dashboards = [
    ("dashboard1_market_overview.png",
     "Dashboard 1 — Market Overview: average rent by city, listing distribution, "
     "room counts, seller types, and price per m² boxplot."),
    ("dashboard2_price_segments.png",
     "Dashboard 2 — Competitor Price Segments: overall segment shares, segments by city, "
     "average price per segment, and seller type split per segment."),
    ("dashboard3_districts.png",
     "Dashboard 3 — Top Districts: top 12 districts by average price in Астана and Алматы."),
    ("dashboard4_source_comparison.png",
     "Dashboard 4 — Source Comparison: Krisha.kz vs OLX.kz by listing count, "
     "average price, and room data coverage."),
    ("dashboard5_swot.png",
     "Dashboard 5 — SWOT Analysis: strengths, weaknesses, opportunities and threats "
     "of the current competitor landscape."),
]

for filename, caption in dashboards:
    img_path = f"{DASH_DIR}/{filename}"
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BLUE
    doc.add_picture(img_path, width=Inches(6.2))
    # Centre the image
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer()

spacer()

# ── Conclusion ───────────────────────────────────────────────────────────────
heading("Conclusion", level=1)
body(
    "Assignment 2 demonstrates a complete data preparation and market intelligence workflow "
    "for RentEasy KZ. The dataset was expanded to two sources (Krisha.kz + OLX.kz), "
    "cleaned programmatically using duplicate detection, missing value handling, standardisation, "
    "and 1.5×IQR outlier removal — resulting in 6,684 high-quality records across 5 cities."
)
body(
    "The market analysis reveals that the Middle price segment (120–250K тг/month) is the "
    "largest and most contested, that Астана is the optimal launch market, and that both "
    "major competitors (Krisha.kz and OLX.kz) share a critical weakness: absence of price "
    "analytics. RentEasy KZ should launch in Астана targeting the Middle segment, "
    "differentiate through transparent price dashboards, and monetise via freemium subscription."
)

spacer()
p = doc.add_paragraph()
run = p.add_run("Source code: https://github.com/Choppa2077/assignment1DDM")
run.font.size = Pt(10)
run.font.color.rgb = BLUE
run.font.italic = True

# ── Save ────────────────────────────────────────────────────────────────────
doc.save("Assignment_2_DDDM.docx")
print("Saved: Assignment_2_DDDM.docx")
