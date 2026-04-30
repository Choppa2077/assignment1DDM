#!/usr/bin/env python3
"""
Assignment 4 - Report Generator (DOCX)
RentEasy KZ — Scientific and Network Analytics
"""
from pathlib import Path
from datetime import datetime

import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WARN] python-docx not installed — DOCX report skipped. Run: pip install python-docx")

DATA_DIR = Path("data")
FIG = Path("output/figures")
NET = Path("output/networks")
TOP = Path("output/topics")
OUT = Path("output/report")
OUT.mkdir(parents=True, exist_ok=True)

# ── helpers ──────────────────────────────────────────────────────────────────

def _hex(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shd(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tcPr.append(shd)


def add_h(doc, text, level=1, color="1E3A8A"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _hex(color)
    return h


def add_para(doc, text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
    p.add_run(text)
    return p


def add_bullet(doc, items):
    for item in items:
        if isinstance(item, tuple):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item[0]).bold = True
            p.add_run(": " + item[1])
        else:
            doc.add_paragraph(item, style="List Bullet")


def add_figure(doc, path: Path, caption: str, width=6.0):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.runs[0] if p.runs else p.add_run(caption)
        r.font.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = _hex("666666")
    else:
        doc.add_paragraph(f"[Figure not generated yet: {path.name}]")


def add_df_table(doc, df: pd.DataFrame, header_fill="1E3A8A"):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        _shd(hdr[i], header_fill)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            v = row[col]
            cells[i].text = str(v) if pd.notna(v) else ""
    return table


# ── main report ──────────────────────────────────────────────────────────────

def build_report():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.2)
        sec.right_margin = Inches(1.2)

    # ── Cover ────────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("RentEasy KZ")
    r.font.size = Pt(30); r.font.bold = True
    r.font.color.rgb = _hex("1E3A8A")

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("Assignment 4 — Scientific and Network Analytics\nfor Innovation Decisions")
    sr.font.size = Pt(16); sr.font.bold = True

    doc.add_paragraph()
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run(
        "Students: Mukan Idrissov, Assylmurat Yeraliev\n"
        "Group: CSE-2501M\n"
        f"Date: {datetime.now().strftime('%B %Y')}\n"
        "Code: https://github.com/Choppa2077/assignment1DDM"
    )
    doc.add_page_break()

    # ── Executive Summary ────────────────────────────────────────────────────
    add_h(doc, "Executive Summary", 1)
    doc.add_paragraph(
        "This report presents Assignment 4 of the RentEasy KZ Data-Driven Decision Making project. "
        "Building on prior rental market analysis (6,684 listings), competitor sentiment analytics, "
        "and Google Trends demand forecasting, this assignment adds a systematic review of peer-reviewed "
        "scientific literature. We collected and analysed over 1,000 publications from OpenAlex on ML/AI "
        "for real estate rental platforms, covering price prediction, NLP review analysis, time-series "
        "forecasting, and recommendation systems."
    )
    doc.add_paragraph(
        "Task 1 built the research dataset. Task 2 produced five bibliometric dashboards, a keyword "
        "co-occurrence network with centrality analysis and community detection, and LDA topic modeling. "
        "Task 3 identifies research gaps, leading research fronts, and concrete recommendations for "
        "incorporating verified scientific solutions into RentEasy KZ."
    )
    doc.add_page_break()

    # ── TASK 1 ───────────────────────────────────────────────────────────────
    add_h(doc, "Task 1 — Scientific Data Collection and Processing", 1)

    add_h(doc, "1.1  Research Topic Formulation", 2)
    doc.add_paragraph(
        "Based on Assignments 1–3, which revealed three core technical needs — "
        "(1) data-driven rental pricing, (2) competitor sentiment monitoring, and "
        "(3) demand forecasting — we defined the following scientific topic:"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run(
        '"Machine Learning and Natural Language Processing for Real Estate Rental Platforms: '
        'Price Prediction, User Sentiment Analysis, and Demand Forecasting"'
    )
    r.font.italic = True; r.font.bold = True

    add_h(doc, "1.2  Search Strategy", 2)
    doc.add_paragraph("Database: OpenAlex (https://openalex.org). "
                       "No API key required; 'mailto' parameter used for polite pool.")
    add_df_table(doc, pd.DataFrame({
        "Search Query": [
            "rental price prediction machine learning regression",
            "real estate sentiment analysis user reviews NLP",
            "housing demand forecasting time series ARIMA deep learning",
            "property recommendation system collaborative filtering",
            "apartment rental platform mobile application UX proptech",
            "real estate price prediction neural network convolutional",
        ],
        "Period": ["2015–2024"] * 5 + ["2018–2024"],
        "Target N": [250, 220, 220, 180, 160, 150],
    }))

    add_h(doc, "1.3  Dataset Overview", 2)
    try:
        df = pd.read_csv(DATA_DIR / "cleaned_publications.csv")
        stats = pd.DataFrame({
            "Metric": [
                "Total unique publications", "Year range",
                "Unique journals", "Unique first authors",
                "Unique countries", "Papers with abstract",
                "Mean citations", "Median citations",
            ],
            "Value": [
                f"{len(df):,}",
                f"{int(df['year'].min())} – {int(df['year'].max())}",
                f"{df['journal'].nunique():,}",
                f"{df['first_author'].nunique():,}",
                f"{df['primary_country'].nunique()}",
                f"{(df['abstract'] != '').sum():,}",
                f"{df['cited_by_count'].mean():.1f}",
                f"{df['cited_by_count'].median():.0f}",
            ]
        })
        add_df_table(doc, stats)
    except Exception:
        doc.add_paragraph("See data/cleaned_publications.csv for statistics.")

    add_h(doc, "1.4  Data Cleaning Process", 2)
    add_bullet(doc, [
        ("Deduplication", "DOI-based first; then case-insensitive title deduplication for records without DOI."),
        ("Abstract reconstruction", "OpenAlex stores abstracts as inverted indexes; reconstructed by sorting word–position pairs."),
        ("Author normalisation", "Unicode NFKC normalisation, whitespace collapse, punctuation trimming."),
        ("Journal normalisation", "Stripped parenthetical qualifiers '(Online)', '(Print)', etc."),
        ("Country extraction", "ISO-3166 2-letter codes mapped to full country names from all author affiliations."),
        ("Missing values", "Empty journal → 'Unknown Journal'; empty country → 'Unknown'; null citations → 0."),
        ("Year filter", "Restricted to 2010–2025 for relevance."),
    ])
    doc.add_page_break()

    # ── TASK 2 ───────────────────────────────────────────────────────────────
    add_h(doc, "Task 2 — Bibliometric Analysis and Network Analytics", 1)

    add_h(doc, "2.1  Publication Dynamics", 2)
    doc.add_paragraph(
        "Figure 1 shows strong growth in publications from 2015 to 2024. The dashed trend line "
        "confirms an accelerating rate of research, especially after 2018, coinciding with the deep "
        "learning surge and increased availability of large rental datasets."
    )
    add_figure(doc, FIG / "fig1_publication_dynamics.png",
               "Figure 1 — Publication dynamics 2015–2024 with linear trend")

    add_h(doc, "2.2  Top-10 Journals", 2)
    doc.add_paragraph(
        "Leading publication venues include journals focused on applied machine learning, smart cities, "
        "urban informatics, and real estate economics — confirming the interdisciplinary character of "
        "this research field. Color encodes average citation impact."
    )
    add_figure(doc, FIG / "fig2_top_journals.png",
               "Figure 2 — Top-10 journals (bar height = count; color = avg citations)")

    add_h(doc, "2.3  Top-10 Authors", 2)
    doc.add_paragraph(
        "The most productive first authors represent potential collaboration targets and benchmark "
        "researchers whose methods are most directly applicable to RentEasy KZ."
    )
    add_figure(doc, FIG / "fig3_top_authors.png",
               "Figure 3 — Top-10 most productive first authors")

    add_h(doc, "2.4  Geographic Distribution", 2)
    doc.add_paragraph(
        "Research is dominated by the United States, China, and Western Europe. Kazakhstan appears "
        "in the corpus, but the near-absence of Central Asian case studies is a significant gap "
        "(see Task 3). The global coverage validates that findings are transferable across markets."
    )
    add_figure(doc, FIG / "fig4_top_countries.png",
               "Figure 4 — Top-10 countries and global choropleth distribution")

    add_h(doc, "2.5  Citation Distribution", 2)
    doc.add_paragraph(
        "Citation counts follow a power-law distribution. The box-plot by year confirms that "
        "older papers (2015–2019) have accumulated more citations. A handful of landmark papers "
        "have 500+ citations, representing the foundational methods in this field."
    )
    add_figure(doc, FIG / "fig5_citation_distribution.png",
               "Figure 5 — Citation distribution (histogram left; box-by-year right)")

    add_h(doc, "2.6  Keyword Co-occurrence Network", 2)
    doc.add_paragraph(
        "We built a keyword co-occurrence network: nodes are keywords, edges connect keywords "
        "that appear together in the same paper. Node size reflects keyword frequency; edge "
        "thickness reflects co-occurrence count. Communities were detected using greedy modularity "
        "maximisation (NetworkX)."
    )
    add_figure(doc, NET / "network_keyword_cooccurrence.png",
               "Figure 6 — Keyword co-occurrence network (left: full; right: community structure)",
               width=6.5)

    try:
        cent = pd.read_csv(NET / "centrality_metrics.csv")
        top10 = cent.head(10)[["keyword", "degree", "betweenness", "pagerank", "frequency"]].copy()
        top10.columns = ["Keyword", "Degree", "Betweenness", "PageRank", "Frequency"]
        doc.add_paragraph("\nTop-10 keywords by PageRank:")
        add_df_table(doc, top10)
    except Exception:
        pass

    doc.add_paragraph(
        "\nKey observations: 'machine learning', 'deep learning', and 'neural network' form the "
        "central methodology hub. 'Sentiment analysis' and 'NLP' bridge the text-mining and real-estate "
        "clusters. 'Time series' and 'forecasting' are tightly connected but peripheral, indicating "
        "forecasting is a specialised sub-field."
    )

    add_h(doc, "2.7  Topic Modeling (LDA)", 2)
    doc.add_paragraph(
        "We applied Latent Dirichlet Allocation (7 topics) to the concatenation of titles, abstracts, "
        "and keywords. The model reveals seven coherent research themes corresponding to the main "
        "technical pillars of modern proptech platforms."
    )
    try:
        topics_df = pd.read_csv(TOP / "lda_topics.csv")
        add_df_table(doc, topics_df[["topic_id", "label", "top_words"]].rename(
            columns={"topic_id": "#", "label": "Topic Label", "top_words": "Top Keywords"}))
    except Exception:
        pass

    add_figure(doc, TOP / "topic_modeling.png",
               "Figure 7 — LDA word clouds per topic + topic trend over time",
               width=6.5)
    add_figure(doc, TOP / "keyword_trends.png",
               "Figure 8 — Keyword trend detection: top keywords 2015–2024")
    doc.add_page_break()

    # ── TASK 3 ───────────────────────────────────────────────────────────────
    add_h(doc, "Task 3 — Research Gaps, Fronts, and Strategic Recommendations", 1)

    add_h(doc, "3.1  Identified Research Gaps", 2)
    add_bullet(doc, [
        ("Central Asian Rental Markets",
         "Virtually no peer-reviewed studies address Kazakhstan, Kyrgyzstan, or other post-Soviet "
         "rental markets. Most datasets are US, Chinese, or Western European. Localised pricing "
         "models and demand patterns remain unstudied."),
        ("Multilingual NLP for CIS Property Platforms",
         "Existing sentiment models are predominantly English. Russian/Kazakh multilingual review "
         "analysis for property apps is absent — directly needed by RentEasy KZ."),
        ("Price Prediction with Sparse / Unverified Data",
         "Most ML pricing models assume large, clean datasets. Methods for working with partially "
         "verified or sparse listings (common in emerging markets) are understudied."),
        ("Listing Authenticity Verification",
         "Fake-listing detection as a specific NLP task in property portals is largely unexplored, "
         "despite being the top complaint against Krisha.kz (Assignment 3)."),
        ("Behavioural Analytics for Rental Search",
         "Combining click-stream / dwell-time behavioural data with property features for renter "
         "intent modelling is rare in academic literature."),
        ("Mobile-first Optimisation in Resource-constrained Environments",
         "Performance optimisation for rental apps on low-end devices / limited bandwidth "
         "(relevant to the KZ market) has no dedicated research stream."),
    ])

    add_h(doc, "3.2  Leading Research Fronts", 2)
    add_bullet(doc, [
        ("Transformer-based Property Valuation",
         "BERT/GPT and multimodal transformers combine listing text, images, and structured data; "
         "currently the highest-cited emerging approach."),
        ("Graph Neural Networks for Spatial Pricing",
         "GNNs capture neighbourhood effects and transport proximity — outperforming tabular models."),
        ("Explainable AI (XAI) for Housing Models",
         "SHAP and LIME are increasingly required by regulators; growing fast after 2021."),
        ("Federated Learning for Privacy-preserving Analytics",
         "Allows training on distributed rental data without centralising sensitive records."),
        ("Computer Vision for Listing Photo Analysis",
         "CNNs / Vision Transformers for automatic room-quality scoring from listing photos."),
    ])

    add_h(doc, "3.3  Scientific Solutions for RentEasy KZ", 2)
    add_df_table(doc, pd.DataFrame({
        "Priority": ["Critical", "Critical", "High", "High", "Medium"],
        "Area": [
            "Rental Price Prediction",
            "Listing Verification",
            "Demand Forecasting",
            "Sentiment Monitoring",
            "Personalisation",
        ],
        "Recommended Approach": [
            "XGBoost / LightGBM on tabular features (proven best on structured RE data); "
            "add GNN layer for spatial effects",
            "Semi-supervised NLP: train on confirmed fake/real listings, apply to new posts; "
            "image consistency check with CV",
            "SARIMA-LSTM hybrid validated against Google Trends data collected in Assignment 3",
            "Fine-tune multilingual BERT on Russian property reviews (addresses language gap)",
            "Matrix Factorisation on user search/click logs; cold-start via content-based KNN",
        ],
    }))

    add_h(doc, "3.4  Proposed Future Research Directions", 2)
    add_bullet(doc, [
        "Publish the 6,684-listing Krisha.kz dataset as an open academic benchmark for Central "
        "Asian rental price prediction — filling the identified data gap.",
        "Build and evaluate a Russian/Kazakh multilingual BERT sentiment model specifically for "
        "property reviews — first such study for the CIS region.",
        "Develop a listing-verification classifier using the Krisha.kz negative keyword patterns "
        "identified in Assignment 3 as training signal.",
        "Benchmark ARIMA vs Prophet vs LSTM-ARIMA hybrid on the Kazakhstan Google Trends series "
        "collected in Assignment 3.",
        "Conduct a GNN-based spatial pricing study using Almaty neighbourhood graph data — "
        "testing whether proximity effects match Western market results.",
    ])

    add_h(doc, "Conclusions", 1)
    doc.add_paragraph(
        "The scientific literature review confirms that all three core technical pillars of RentEasy KZ "
        "— price prediction, sentiment monitoring, and demand forecasting — are supported by mature, "
        "high-quality research. The identified gaps (Central Asian data, Russian NLP, verification "
        "systems) simultaneously represent genuine scientific contributions and sustainable competitive "
        "advantages. The keyword network and LDA analysis show the field is converging on transformer "
        "and GNN architectures, pointing to the long-term technology investment direction."
    )

    add_h(doc, "Data and Code", 1)
    add_bullet(doc, [
        "data/raw_publications.json — raw OpenAlex API response",
        "data/cleaned_publications.csv — final analytical dataset",
        "output/figures/ — 5 interactive dashboards (HTML + PNG)",
        "output/networks/ — co-occurrence network + centrality CSV",
        "output/topics/ — LDA topics CSV + visualisations",
        "GitHub: https://github.com/Choppa2077/assignment1DDM",
    ])

    path = OUT / "RentEasy_KZ_Assignment4_Report.docx"
    doc.save(str(path))
    print(f"\nDOCX report saved → {path}")


def main():
    if not HAS_DOCX:
        print("Install python-docx and re-run: pip install python-docx")
        return
    build_report()


if __name__ == "__main__":
    main()
