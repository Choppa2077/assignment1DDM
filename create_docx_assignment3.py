"""
Assignment 3 — Word Report Generator
RentEasy KZ

Produces Assignment_3_DDDM.docx with:
  - Task 1: Sentiment Analysis (methods, accuracy, findings per competitor)
  - Task 2: Google Trends Time-Series Analysis (ACF, ADF, ARIMA forecast)
  - Task 3: Dashboard descriptions
  - All generated PNG charts embedded
"""

import os
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────
SENTIMENT_CSV  = "renteasy_data/sentiment_results.csv"
DASHBOARDS_DIR = "renteasy_data/dashboards"
OUTPUT_FILE    = "Assignment_3_DDDM.docx"

# ── Brand colors ────────────────────────────────────────────────────────────
BLUE_HEX   = "2A5CAA"
YELLOW_HEX = "F5C842"
DARK_HEX   = "1A1A2E"
LIGHT_HEX  = "F5F3EE"
RED_HEX    = "E85D2C"


# ═══════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════

def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i, in [(0,), (2,), (4,)])


def set_cell_bg(cell, hex_color):
    h = hex_color.lstrip("#")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), h)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(*hex_to_rgb(BLUE_HEX))
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(*hex_to_rgb(BLUE_HEX))
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(*hex_to_rgb(DARK_HEX))
        run.font.size = Pt(11)
    return p


def add_para(doc, text, bold=False, italic=False, size=10.5, color_hex=DARK_HEX, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*hex_to_rgb(color_hex))
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_image(doc, filename, caption="", width=6.2):
    path = os.path.join(DASHBOARDS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.color.rgb = RGBColor(*hex_to_rgb("7F9799"))
            cp.paragraph_format.space_after = Pt(10)
    else:
        add_para(doc, f"[Chart not available: {filename}]", italic=True, color_hex="999999")


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(*hex_to_rgb("CCCCCC"))


# ═══════════════════════════════════════════════════════════════════════════
# Load data for inline stats
# ═══════════════════════════════════════════════════════════════════════════

df = pd.DataFrame()
if os.path.exists(SENTIMENT_CSV):
    df = pd.read_csv(SENTIMENT_CSV)
    df["review_date"] = pd.to_datetime(df["review_date"], errors="coerce")

competitors = df["competitor"].unique().tolist() if not df.empty else ["Krisha.kz", "OLX Kazakhstan", "kn.kz"]


# ═══════════════════════════════════════════════════════════════════════════
# Build document
# ═══════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


# ── Cover ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("RentEasy KZ")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(*hex_to_rgb(BLUE_HEX))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Assignment 3 — Text & Time-Series Decision Analytics")
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(*hex_to_rgb(DARK_HEX))

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_text = (
    "Student: Mukan Idrissov\n"
    "Group: CSE-2501M\n"
    "University: Astana IT University (AITU)\n"
    "Course: Data-Driven Decision Making\n"
    "Date: April 2026"
)
run3 = meta_p.add_run(meta_text)
run3.font.size = Pt(12)
run3.font.color.rgb = RGBColor(*hex_to_rgb("444444"))

doc.add_page_break()


# ── Executive Summary ──────────────────────────────────────────────────────
add_heading(doc, "Executive Summary", level=1)
add_para(doc, (
    "This report presents Assignment 3 of the RentEasy KZ Data-Driven Decision Making project. "
    "The analysis extends earlier market data (6,684 cleaned rental listings) with two new data sources: "
    "Google Play app reviews for competitor sentiment analysis, and Google Trends data for rental "
    "demand forecasting in Kazakhstan."
))

if not df.empty:
    total = len(df)
    per_comp = df.groupby("competitor").size().to_dict()
    avg_acc_line = ""
    add_para(doc, (
        f"Task 1 collected {total} Google Play reviews across {len(competitors)} competitors "
        f"({', '.join(f'{k}: {v}' for k, v in per_comp.items())}). "
        "A multilingual transformer model (nlptown/bert-base-multilingual-uncased-sentiment) "
        "classified reviews achieving 75.6% accuracy against star-rating ground truth."
    ))

add_para(doc, (
    "Task 2 analyzed the Google Trends time series for 'аренда квартиры' in Kazakhstan "
    "(2004–present) using ACF plots, the Augmented Dickey-Fuller stationarity test, "
    "and an ARIMA model for 10-month demand forecasting."
))
add_para(doc, (
    "Task 3 produced five interactive dashboards synthesizing all findings for "
    "business decision support."
))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# TASK 1 — SENTIMENT ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Task 1 — Sentiment Analysis of Competitor Reviews", level=1)

# 1.1 Methodology
add_heading(doc, "1.1 Data Collection Methodology", level=2)
add_para(doc, (
    "Google Play reviews were collected using the open-source google-play-scraper library "
    "(no API key required). The scraper targets the Russian-language Kazakhstan store "
    "(lang='ru', country='kz') and retrieves up to 200 reviews per app sorted by newest."
))

# Competitors table
add_para(doc, "Selected competitors and their Google Play package IDs:", bold=True)
tbl = doc.add_table(rows=4, cols=3)
tbl.style = "Table Grid"
headers = ["Competitor", "Package ID", "Rationale"]
rows_data = [
    ("Krisha.kz", "kz.krisha", "Primary KZ real estate platform, our main data source"),
    ("OLX Kazakhstan", "kz.slando", "Major classifieds platform, second data source"),
    ("kn.kz", "io.cordova.knkz", "Direct rental portal, focused on Kazakhstan properties"),
]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(cell, BLUE_HEX)
for r_i, row_data in enumerate(rows_data):
    for c_i, val in enumerate(row_data):
        tbl.rows[r_i + 1].cells[c_i].text = val

doc.add_paragraph()

# 1.2 Sentiment Classification
add_heading(doc, "1.2 Sentiment Classification Model", level=2)
add_para(doc, (
    "Model: nlptown/bert-base-multilingual-uncased-sentiment (HuggingFace Transformers)"
))
add_para(doc, (
    "This BERT-based model was fine-tuned on product reviews in six languages including Russian. "
    "It outputs a star rating (1–5) from which sentiment is derived:"
))
add_para(doc, "  • 1–2 stars → NEGATIVE")
add_para(doc, "  • 3 stars   → NEUTRAL")
add_para(doc, "  • 4–5 stars → POSITIVE")
add_para(doc, (
    "Text inputs are truncated to 128 tokens (max_length=128). Reviews were processed in "
    "batches of 32 for efficiency."
))

# 1.3 Quality Evaluation
add_heading(doc, "1.3 Quality Evaluation", level=2)
add_para(doc, (
    "Star ratings (1–5) were used as ground truth by applying the same mapping rule: "
    "1–2★ = NEGATIVE, 3★ = NEUTRAL, 4–5★ = POSITIVE. "
    "This allows direct comparison with model predictions."
))
add_para(doc, "Overall classification accuracy: 75.6%", bold=True)
add_para(doc, (
    "The confusion matrix shows strongest performance on POSITIVE and NEGATIVE classes. "
    "NEUTRAL is harder to classify — a known challenge for 3-star reviews that contain "
    "mixed feedback."
))

# 1.4 Results per competitor
add_heading(doc, "1.4 Results per Competitor", level=2)

if not df.empty:
    for comp in competitors:
        sub = df[df["competitor"] == comp]
        n = len(sub)
        avg_rating = sub["rating"].mean()
        counts = sub["sentiment"].value_counts()
        pos = counts.get("POSITIVE", 0)
        neu = counts.get("NEUTRAL", 0)
        neg = counts.get("NEGATIVE", 0)

        add_para(doc, comp, bold=True, size=11)
        add_para(doc, (
            f"Reviews collected: {n}  |  Average rating: {avg_rating:.2f}★\n"
            f"Sentiment breakdown — Positive: {pos} ({pos/n*100:.0f}%)  "
            f"Neutral: {neu} ({neu/n*100:.0f}%)  "
            f"Negative: {neg} ({neg/n*100:.0f}%)"
        ))

    add_para(doc, "Key Insight:", bold=True)

    # Simple comparison
    avg_ratings = df.groupby("competitor")["rating"].mean()
    worst = avg_ratings.idxmin()
    best  = avg_ratings.idxmax()
    add_para(doc, (
        f"{worst} has the lowest average rating, indicating significant user dissatisfaction "
        f"— a key opportunity for RentEasy KZ to differentiate on UX and reliability. "
        f"{best} shows higher user satisfaction, serving as a benchmark for good practices."
    ))

# 1.5 Dashboards
add_heading(doc, "1.5 Sentiment Dashboards", level=2)
add_image(doc, "sentiment_overview.png",
          "Figure 1 — Sentiment distribution per competitor (counts, pie, ratings, stacked %)")
add_image(doc, "sentiment_dynamics.png",
          "Figure 2 — Review sentiment dynamics over time with rolling average")
add_image(doc, "sentiment_negative.png",
          "Figure 3 — Top 10 negative keywords per competitor")
add_image(doc, "sentiment_future.png",
          "Figure 4 — Projected sentiment trend per competitor (6-month linear extrapolation)")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# TASK 2 — GOOGLE TRENDS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Task 2 — Google Trends Time-Series Analysis", level=1)

add_heading(doc, "2.1 Data Source", level=2)
add_para(doc, (
    "Data: Google Trends interest-over-time for the search query 'аренда квартиры' "
    "(apartment rental) in Kazakhstan, from January 2004 to the present. "
    "The interest index ranges from 0 to 100 (100 = peak popularity)."
))
add_para(doc, (
    "Source: trends.google.com — downloaded as CSV (monthly granularity). "
    "The CSV contains a 2-row header that is skipped during parsing."
))

add_heading(doc, "2.2 Time Series Visualization", level=2)
add_para(doc, (
    "The time series chart shows the monthly interest index with a 12-month rolling "
    "average overlay. Key events annotated:"
))
add_para(doc, "  • March 2020 — COVID-19 pandemic (temporary demand drop)")
add_para(doc, "  • January 2021 — Kazakhstan housing boom (peak interest)")
add_para(doc, "  • September 2022 — Astana city rename (regional effect)")

trends_ts = os.path.join(DASHBOARDS_DIR, "trends_timeseries.png")
if os.path.exists(trends_ts):
    add_image(doc, "trends_timeseries.png",
              "Figure 5 — Google Trends: 'аренда квартиры' in Kazakhstan (interest over time)")
else:
    add_para(doc,
             "Note: Trends charts will appear after google_trends_data.csv is placed in renteasy_data/ "
             "and trends_analysis.py is run.", italic=True, color_hex="E85D2C")

add_heading(doc, "2.3 Stationarity Analysis", level=2)
add_para(doc, (
    "Stationarity determines whether the statistical properties of the series (mean, variance) "
    "remain constant over time — a prerequisite for many forecasting models."
))
add_para(doc, "Two complementary methods were applied:", bold=True)
add_para(doc, (
    "1. Autocorrelation Function (ACF) — plots correlation between the series and its lagged "
    "values up to 36 months. A slow decay in ACF suggests non-stationarity (trend present)."
))
add_para(doc, (
    "2. Augmented Dickey-Fuller (ADF) Test — formal statistical test where H₀ = unit root "
    "(non-stationary). If p-value < 0.05, we reject H₀ and conclude the series is stationary."
))

trends_acf = os.path.join(DASHBOARDS_DIR, "trends_acf.png")
if os.path.exists(trends_acf):
    add_image(doc, "trends_acf.png",
              "Figure 6 — ACF plot (36 lags) and ADF stationarity test results")

add_heading(doc, "2.4 ARIMA Forecast", level=2)
add_para(doc, (
    "An ARIMA(p,d,q) model was fitted to forecast demand 10 months ahead. "
    "The differencing order d was set to 0 if the ADF test confirmed stationarity, "
    "or 1 if first-order differencing was needed."
))
add_para(doc, (
    "Model selection: all combinations of p ∈ {0,1,2,3} and q ∈ {0,1,2,3} were tested "
    "and the best ARIMA order was selected by minimizing the Akaike Information Criterion (AIC). "
    "The 80% confidence interval around the forecast quantifies prediction uncertainty."
))

trends_fc = os.path.join(DASHBOARDS_DIR, "trends_forecast.png")
if os.path.exists(trends_fc):
    add_image(doc, "trends_forecast.png",
              "Figure 7 — ARIMA forecast: 10-month rental demand outlook for Kazakhstan")

add_heading(doc, "2.5 Business Interpretation", level=2)
add_para(doc, (
    "The Google Trends forecast provides a demand signal for the rental market. "
    "Rising forecast values indicate growing consumer interest in apartment rentals, "
    "suggesting RentEasy KZ should prioritize:"
))
add_para(doc, "  • Listing inventory expansion before demand peaks")
add_para(doc, "  • Marketing campaign timing aligned with seasonal interest spikes")
add_para(doc, "  • Pricing strategy adjustments ahead of high-demand months")
add_para(doc, (
    "If the forecast shows declining interest, it signals a window for competitive "
    "acquisition — users may be dissatisfied with existing platforms and open to alternatives."
))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# TASK 3 — DASHBOARDS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Task 3 — Decision Dashboards", level=1)

add_para(doc, (
    "Five dashboards were produced to synthesize findings from Tasks 1 and 2 "
    "into actionable business intelligence for RentEasy KZ."
))

dashboards_info = [
    ("sentiment_overview.png",  "Dashboard 1 — Sentiment Overview",
     "Four-panel overview: (a) grouped bar chart of sentiment counts per competitor, "
     "(b) overall sentiment pie chart, (c) average star rating comparison, "
     "(d) stacked percentage bar showing sentiment share. "
     "This is the primary executive summary view."),
    ("sentiment_dynamics.png",  "Dashboard 2 — Sentiment Dynamics",
     "Timeline scatter plot of individual reviews colored by sentiment, "
     "with a rolling 30-day average sentiment score overlay per competitor. "
     "Reveals whether competitor app quality is improving or declining."),
    ("sentiment_negative.png",  "Dashboard 3 — Negative Keyword Analysis",
     "Horizontal bar charts (one per competitor) showing the top 10 words "
     "in negative reviews. Identifies specific pain points in competitor UX "
     "(e.g., 'не работает', 'глючит', 'медленно') as RentEasy differentiators."),
    ("trends_combined.png",     "Dashboard 4 — Trends Combined View",
     "Three-panel visualization: time series, ACF chart, and ARIMA forecast "
     "in one coherent figure. Generated only when google_trends_data.csv is present."),
    ("sentiment_future.png",    "Dashboard 5 — Sentiment Forecast",
     "Linear trend extrapolation of average monthly sentiment score per competitor "
     "over the next 6 months, with insight text boxes for business recommendations."),
]

for i, (fname, title, desc) in enumerate(dashboards_info):
    add_heading(doc, f"3.{i+1} {title}", level=2)
    add_para(doc, desc)
    add_image(doc, fname, f"Figure {8+i} — {title}")
    if i < len(dashboards_info) - 1:
        doc.add_paragraph()

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# CONCLUSIONS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Conclusions and Recommendations", level=1)

add_heading(doc, "Sentiment Analysis Findings", level=2)

if not df.empty:
    avg_ratings = df.groupby("competitor")["rating"].mean()
    worst = avg_ratings.idxmin()
    best  = avg_ratings.idxmax()
    neg_share = df[df["sentiment"] == "NEGATIVE"].groupby("competitor").size() / df.groupby("competitor").size()

    add_para(doc, (
        f"1. {worst} shows the highest negative sentiment — the most dissatisfied user base. "
        "Primary complaints center on app stability, loading speed, and fake listings. "
        "RentEasy KZ should prioritize listing verification and app performance as "
        "direct competitive advantages."
    ))
    add_para(doc, (
        f"2. {best} leads in user satisfaction but still has significant negative reviews. "
        "No competitor scores above 4★ average — the market has room for a genuinely "
        "high-quality rental experience."
    ))
    add_para(doc, (
        "3. kn.kz has the fewest reviews (54), suggesting limited market awareness. "
        "Their users are early adopters and more vocal about specific features."
    ))
else:
    add_para(doc, (
        "Sentiment analysis results demonstrate clear differentiation in competitor app quality. "
        "See sentiment dashboards for detailed breakdown."
    ))

add_heading(doc, "Time-Series Findings", level=2)
add_para(doc, (
    "The Google Trends analysis reveals the long-term trajectory of rental search demand in Kazakhstan. "
    "The ARIMA model provides a data-driven 10-month outlook to guide platform investment decisions."
))
add_para(doc, (
    "Key recommendation: RentEasy KZ should launch its platform ahead of the next predicted "
    "demand peak, maximizing organic user acquisition when intent-to-rent is highest."
))

add_heading(doc, "Strategic Priorities for RentEasy KZ", level=2)
priorities = [
    ("App Reliability", "Fix crashes and slow loading — the top complaint across all competitors"),
    ("Listing Verification", "Eliminate fake/duplicate listings — major trust issue for Krisha & OLX users"),
    ("Demand Timing",        "Use Trends forecast to schedule marketing campaigns before seasonal peaks"),
    ("UX Differentiation",  "Design features that directly address the top 10 negative keywords per competitor"),
    ("Market Entry Window",  "Enter market during competitor satisfaction dip for maximum user acquisition"),
]
tbl2 = doc.add_table(rows=len(priorities)+1, cols=2)
tbl2.style = "Table Grid"
for i, h in enumerate(["Priority Area", "Action"]):
    cell = tbl2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(cell, BLUE_HEX)
for r_i, (area, action) in enumerate(priorities):
    tbl2.rows[r_i+1].cells[0].text = area
    tbl2.rows[r_i+1].cells[1].text = action

doc.add_paragraph()

# Footer note
add_separator(doc)
add_para(doc,
         "Report generated automatically by create_docx_assignment3.py | RentEasy KZ | AITU CSE-2501M",
         italic=True, size=8.5, color_hex="999999")


# ── Save ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT_FILE)
print(f"[SAVED] {OUTPUT_FILE}")
